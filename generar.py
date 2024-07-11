import pandas as pd
from openpyxl import Workbook
from docx import Document

# Datos de la tabla
data = [
    ["No.", "ACTIVIDADES", "Comité Colaborador del Proyecto", "Comisión del Cliente", "Administrador del Modelo", "Ingeniero de Sistemas", "Técnico de Sistemas", "Supervisor", "Electricistas", "Personal del Modelo"],
    [1, "Contratación del Proyecto", "R", "S", "", "", "", "", "", ""],
    [2, "Definir alcance del proyecto", "R", "", "S", "", "", "", "", ""],
    [3, "Creación del cronograma de actividades", "R", "", "S", "", "", "", "", ""],
    [4, "Asignación de recursos", "R", "", "S", "", "", "", "", ""],
    [5, "Contratación del Administrador del modelo", "R", "S", "", "", "", "", "", ""],
    [6, "Difusión de la propuesta en la institución", "R", "", "S", "", "", "", "", ""],
    [7, "Contratación del personal técnico", "R", "", "S", "", "", "", "", ""],
    [8, "Preparación del ambiente de pruebas", "", "", "S", "R", "", "", "", ""],
    [9, "Adquisición de equipo y materiales", "R", "", "S", "", "", "", "", ""],
    [10, "Capacitación del personal", "R", "", "S", "", "", "", "", ""],
    [11, "Entrega preliminar del modelo", "", "", "S", "R", "", "", "", ""],
    [12, "Pruebas preliminares", "", "", "S", "R", "", "", "", ""],
    [13, "Retroalimentación de las pruebas preliminares", "", "", "S", "R", "", "", "", ""],
    [14, "Liberación de recursos", "R", "", "S", "", "", "", "", ""],
    [15, "Inicio de la marcha blanca", "", "", "S", "R", "", "", "", ""],
    [16, "Evaluación final", "", "", "S", "R", "", "", "", ""]
]

# Crear un DataFrame
df = pd.DataFrame(data[1:], columns=data[0])

# Guardar el DataFrame en un archivo Excel
excel_file = "table.xlsx"
df.to_excel(excel_file, index=False)

# Crear un nuevo documento de Word
doc = Document()

# Agregar un título al documento
doc.add_heading('Tabla de Actividades del Plan de Implantación', level=1)

# Agregar una tabla al documento de Word
table = doc.add_table(rows=1, cols=len(data[0]))

# Agregar los encabezados de la tabla
hdr_cells = table.rows[0].cells
for i, column_name in enumerate(data[0]):
    hdr_cells[i].text = column_name

# Agregar los datos a la tabla
for row in data[1:]:
    row_cells = table.add_row().cells
    for i, cell in enumerate(row):
        row_cells[i].text = str(cell)

# Guardar el documento de Word
word_file = "table.docx"
doc.save(word_file)

print(f"Tabla guardada en {excel_file} y {word_file}")
